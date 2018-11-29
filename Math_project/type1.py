from py2neo import Graph
import re
graph = Graph("bolt://localhost:7687",username="neo4j",password="13820541017")
import docx
from PIL import Image
import requests

file = docx.Document()
def get_data_from_database(cypher):
    data = graph.run(cypher).to_data_frame()
    data = data.values.tolist()
    return data

def main():
    cpyher = ("match (n:read2) return n.a_reading,n.b_answer")
    data = get_data_from_database(cpyher)
    for i in range (len(data)+1):
        # print(data[0][i])
        string = data[0][i]
        pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',re.S)
        url = re.findall(pattern,string)
        file.add_paragraph(data[0][i])
        if len(url)>0:
            img = url[0]
            img_name = img.split('/')[-1]
            with open(img_name, 'wb')as f:
                response = requests.get(img).content
                f.write(response)
                f.close()
            del url[0]
            file.add_picture(img_name)


    file.save("D:\\English.docx")


if __name__ == "__main__":
    main()
