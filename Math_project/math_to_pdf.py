from py2neo import Graph
import docx
import os
from win32com import client
from selenium import webdriver
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait

#定义word文档，并生成文件
file = docx.Document()
file.add_paragraph("******************************  北京雄恩教育研发中心研发  *******************************"+"\n")
file.add_paragraph("                                                                               函数 ")
#定义链接数据库
graph = Graph('bolt://localhost:7687',username="neo4j",password="13820541017")

#定义接口，使用户能够自己定义文件的命名等
print("请输入您希望保存的文档名字：",end="")
file_name = input()

#定义网站驱动
browser = webdriver.Chrome()
browser.maximize_window()

#定义函数获取数据
def get_data_from_database(cypher):
    data = graph.run(cypher).to_data_frame()
    data = data.values.tolist()
    k = 1
    for i in range (len(data)):
        if data[i][1] == None:
            print((data[i][0]))
            file.add_paragraph(data[i][0])

        else:
            print(str(data[i][1])+" "+str(data[i][0]))
            file.add_paragraph(str(k)+"、"+str(data[i][1])+"："+str(data[i][0]))
            k+=1

def doc_to_pdf(doc_name, pdf_name):

    try:
        word = client.DispatchEx("Word.Application")
        if os.path.exists(pdf_name):
              os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name, ReadOnly=1)
        worddoc.SaveAs(pdf_name, FileFormat=17)
        worddoc.Close()
        return pdf_name
    except:
        return 1

def main():
    cypher = "match n=(core:new_core_elements14{define:'值域'})<-[*..3]-(x) return x.define,x.content"
    get_data_from_database(cypher)

    file.save("D:\\python_web_server" + "\\" + str(file_name) + ".docx")
    doc_name = "D:\\python_web_server" + "\\" + str(file_name) + ".docx"
    ftp_name = "D:\\python_web_server\\" + str(file_name) + ".pdf"

    doc_to_pdf(doc_name, ftp_name)
    browser.get("http://localhost:8081/" + str(file_name) + ".pdf")

if __name__ == "__main__":
    main()